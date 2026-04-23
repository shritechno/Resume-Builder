from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt

from app.models.schemas import ResumeDownloadRequest


def build_resume_doc(template_name: str, template_slug: str, payload: ResumeDownloadRequest) -> bytes:
    document = Document()
    style_name = "Normal"
    if style_name in document.styles:
        document.styles[style_name].font.name = "Calibri"
        document.styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        document.styles[style_name].font.size = Pt(10.5)

    if template_slug == "modern-professional":
        build_modern_professional(document, payload)
    elif template_slug == "executive-classic":
        build_executive_classic(document, payload)
    else:
        build_simple_ats(document, payload)

    footer = document.add_paragraph()
    footer_run = footer.add_run(f"Template: {template_name}")
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_simple_ats(document: Document, payload: ResumeDownloadRequest) -> None:
    title = document.add_paragraph()
    title_run = title.add_run(payload.full_name)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(31, 41, 55)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(" | ".join(filter(None, [payload.email, payload.phone, payload.location])))
    subtitle_run.font.color.rgb = RGBColor(75, 85, 99)

    add_section(document, "Professional Summary", [payload.summary], RGBColor(37, 99, 235))
    add_section(document, "Skills", payload.skills, RGBColor(37, 99, 235))
    add_section(document, "Experience", payload.experience, RGBColor(37, 99, 235))
    add_section(document, "Education", payload.education, RGBColor(37, 99, 235))
    if payload.projects:
        add_section(document, "Projects", payload.projects, RGBColor(37, 99, 235))


def build_modern_professional(document: Document, payload: ResumeDownloadRequest) -> None:
    header_table = document.add_table(rows=1, cols=1)
    set_cell_background(header_table.cell(0, 0), "0F766E")
    header_paragraph = header_table.cell(0, 0).paragraphs[0]
    name_run = header_paragraph.add_run(payload.full_name)
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(255, 255, 255)

    contact = header_table.cell(0, 0).add_paragraph(
        " | ".join(filter(None, [payload.email, payload.phone, payload.location]))
    )
    for run in contact.runs:
        run.font.color.rgb = RGBColor(229, 255, 250)
        run.font.size = Pt(10)

    document.add_paragraph("")
    content_table = document.add_table(rows=1, cols=2)
    content_table.autofit = True
    left_cell = content_table.cell(0, 0)
    right_cell = content_table.cell(0, 1)
    set_cell_background(right_cell, "F0FDFA")

    add_cell_section(left_cell, "Professional Summary", [payload.summary], RGBColor(15, 118, 110))
    add_cell_section(left_cell, "Experience", payload.experience, RGBColor(15, 118, 110))
    add_cell_section(left_cell, "Projects", payload.projects, RGBColor(15, 118, 110))

    add_cell_section(right_cell, "Skills", payload.skills, RGBColor(13, 148, 136))
    add_cell_section(right_cell, "Education", payload.education, RGBColor(13, 148, 136))


def build_executive_classic(document: Document, payload: ResumeDownloadRequest) -> None:
    header_table = document.add_table(rows=1, cols=1)
    set_cell_background(header_table.cell(0, 0), "1F2937")
    title = header_table.cell(0, 0).paragraphs[0]
    title_run = title.add_run(payload.full_name)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(255, 255, 255)

    meta = header_table.cell(0, 0).add_paragraph(" | ".join(filter(None, [payload.email, payload.phone, payload.location])))
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(209, 213, 219)

    document.add_paragraph("")
    add_section(document, "Executive Summary", [payload.summary], RGBColor(31, 41, 55))
    add_section(document, "Core Skills", payload.skills, RGBColor(31, 41, 55))
    add_section(document, "Professional Experience", payload.experience, RGBColor(31, 41, 55))
    add_section(document, "Education", payload.education, RGBColor(31, 41, 55))
    if payload.projects:
        add_section(document, "Key Projects", payload.projects, RGBColor(31, 41, 55))


def add_section(document: Document, heading: str, items: list[str], heading_color: RGBColor) -> None:
    filtered_items = [item.strip() for item in items if item and item.strip()]
    if not filtered_items:
        return

    heading_paragraph = document.add_paragraph()
    heading_run = heading_paragraph.add_run(heading.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(11)
    heading_run.font.color.rgb = heading_color

    for item in filtered_items:
        if len(filtered_items) == 1 and "\n" not in item and len(item) < 120 and "summary" in heading.lower():
            document.add_paragraph(item)
        else:
            document.add_paragraph(item, style="List Bullet")


def add_cell_section(cell, heading: str, items: list[str], heading_color: RGBColor) -> None:
    filtered_items = [item.strip() for item in items if item and item.strip()]
    if not filtered_items:
        return

    heading_p = cell.add_paragraph()
    heading_run = heading_p.add_run(heading.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(10.5)
    heading_run.font.color.rgb = heading_color

    for item in filtered_items:
        p = cell.add_paragraph(item)
        if len(filtered_items) > 1:
            p.style = "List Bullet"


def set_cell_background(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)
